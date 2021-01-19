import json
from datetime import date
from io import BytesIO

from django.contrib import messages
from django.http import HttpResponse, Http404
from django.shortcuts import get_object_or_404, redirect
from django.urls import reverse
from django.views.generic import DetailView, CreateView, DeleteView
from django.views.generic.base import View, TemplateView
from docx import *

from webapp.forms import ProgramForm
from webapp.models import Program, PROGRAM_STATUS_CLOSED, PROGRAM_STATUS_OPEN, Child, ProgramSkill, \
    GOAL_STATUS_OPEN, GOAL_STATUS_CLOSED, ProrgamSkillGoal, Skill, SkillLevel, Session


class ProgramDetailView(DetailView):
    template_name = 'program/program_detail_view.html'
    model = Program
    context_object_name = 'programs'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        program = get_object_or_404(Program, pk=self.kwargs.get('pk'))
        goals = ProrgamSkillGoal.objects.filter(skill__program=program)
        goal_open = goals.filter(status=GOAL_STATUS_OPEN)
        pr_skill = program.program_skill.all().order_by('level')
        if not goal_open:
            program.status = PROGRAM_STATUS_CLOSED
        else:
            program.status = PROGRAM_STATUS_OPEN
        program.save()
        context['skills'] = pr_skill
        return context


class ProgramCreateView(CreateView):
    model = Program
    template_name = 'program/program_create.html'
    form_class = ProgramForm


    def form_valid(self, form):
        child = get_object_or_404(Child, pk=self.kwargs.get('pk'))
        program = form.save(commit=False)
        program.child = child
        form.instance.author = self.request.user
        program.save()
        return super().form_valid(form)

    def get_success_url(self):
        program = get_object_or_404(Program, pk=self.object.pk)
        return reverse('webapp:update_program', kwargs={'pk': program.pk})




class UpdateProgram(TemplateView):
    template_name = 'program/program_add_skill.html'
    model = Program

    def get_context_data(self, **kwargs):
        program = get_object_or_404(Program, pk=self.kwargs.get('pk'))
        child = program.child
        test = child.test.last()
        context = super().get_context_data(**kwargs)
        if test:
            category_cod = self.request.GET.get('ABC')
            context['skills'] = Skill.objects.filter(category__code=category_cod)
            context['child_pk'] = test.child.pk
            context['program'] = program
            context['test'] = test
            context['ABC'] = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M',
                              'N', 'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z']

        else:
            context['error'] = 'Нет ни одного проведенного теста'
        return context

    def post(self, request, *args, **kwargs):
        program = get_object_or_404(Program, pk=self.kwargs.get('pk'))
        data = json.loads(request.body)
        skill_lvl = SkillLevel.objects.get(pk=data['id'])
        prorgam_skill = ProgramSkill()
        if data['add_creteria']:
            prorgam_skill.add_creteria = data['add_creteria']
        prorgam_skill.level = skill_lvl
        prorgam_skill.program = program
        prorgam_skill.save()
        goals = [g for g in data['goals'] if g]
        if goals:
            for g in goals:
                goal = ProrgamSkillGoal()
                goal.skill = prorgam_skill
                goal.goal = g
                goal.save()
        else:
            goal = ProrgamSkillGoal()
            goal.skill = prorgam_skill
            goal.save()

        return redirect('webapp:update_program', pk=program.pk)


class ProgrmDelete(DeleteView):
    template_name = 'program/delete_program.html'
    model = Program

    def get_success_url(self):
        return reverse('webapp:child_view', kwargs={'pk': self.object.child.pk})


class DeleteGoalView(DeleteView):
    # template_name = 'program/delete_goal.html'
    model = ProrgamSkillGoal

    def get(self, request, *args, **kwargs):
        prskill = get_object_or_404(ProrgamSkillGoal, pk=self.kwargs.get('pk'))
        session = Session.objects.filter(program=prskill.skill.program)
        if session:
            messages.error(self.request, 'Нельзя удалить цель ведуться ссесии')
            return redirect('webapp:program_detail', prskill.skill.program.pk)
        else:
            return self.delete(request, *args, **kwargs)



    def get_success_url(self):
        return reverse('webapp:program_detail', kwargs={'pk': self.object.skill.program.pk})


class DeleteCreteriaView(DeleteView):
    template_name = 'program/program_detail_view.html'
    model = ProgramSkill

    def get(self, request, *args, **kwargs):
        prskill = get_object_or_404(ProgramSkill, pk=self.kwargs.get('pk'))
        session = Session.objects.filter(program=prskill.program)
        if session:
            messages.error(self.request, 'Нельзя удалить навык  ведуться ссесии')
            return redirect('webapp:program_detail', prskill.program.pk)
        else:
            return self.delete(request, *args, **kwargs)

    def get_success_url(self):
        return reverse('webapp:program_detail', kwargs={'pk': self.object.program.pk})


class DeleteAddCreteriaView(View):
    def get(self, request, *args, **kwargs):
        prskill = get_object_or_404(ProgramSkill, pk=self.kwargs.get('pk'))
        session = Session.objects.filter(program=prskill.program)
        if session:
            messages.error(self.request, 'По программе ведуться ссесии')
            return redirect('webapp:program_detail', prskill.program.pk)
        else:
            prskill.add_creteria = None
            self.prskill = prskill
            prskill.save()
            return self.get_success_url()

    def get_success_url(self):
        return redirect("webapp:program_detail", self.prskill.program.pk)

class RemoveProgramView(TemplateView):
    model = Program

    def get(self, request, *args, **kwargs):
        skill_lvl = SkillLevel.objects.get(pk=self.kwargs.get('s_pk'))
        p_pk = self.kwargs.get('p_pk')
        program_skill = skill_lvl.program_skill.all().filter(program__pk=p_pk)[0]
        session = Session.objects.filter(program=program_skill.program)
        if session:
            responseData = {
                # 'criteria': skill_lvl.criteria,
                'error': 'По данному навыку введетсся ссесия'
            }
            return HttpResponse(json.dumps(responseData), content_type="application/json")


        else:
            goals = program_skill.goal.all()
            g = []
            for goal in goals:
                g.append(goal.goal)
            print(goals)
            responseData = {
                # 'criteria': skill_lvl.criteria,
                'add_creteria': program_skill.add_creteria,
                'goals': g or []
            }
            return HttpResponse(json.dumps(responseData), content_type="application/json")

    def post(self, request, *args, **kwargs):

        skill_lvl = SkillLevel.objects.get(pk=self.kwargs.get('s_pk'))
        p_pk = self.kwargs.get('p_pk')
        program_skill = skill_lvl.program_skill.all().filter(program__pk=p_pk)[0]
        program_skill.goal.all().delete()

        program = get_object_or_404(Program, pk=self.kwargs.get('p_pk'))
        data = json.loads(request.body)
        skill_lvl = SkillLevel.objects.get(pk=data['id'])

        if data['add_creteria'] and not program_skill:
            program_skill = ProgramSkill()
            program_skill.add_creteria = data['add_creteria']
            program_skill.level = skill_lvl
            program_skill.program = program
        elif not data['add_creteria'] and program_skill:
            program_skill.add_creteria = None
        else:
            program_skill.add_creteria = data['add_creteria']
        program_skill.save()
        goals = data['goals']
        if goals:
            for g in goals:
                if g:
                    goal = ProrgamSkillGoal()
                    goal.skill = program_skill
                    goal.goal = g
                    goal.save()
        else:
            goal = ProrgamSkillGoal()
            goal.save()

        return redirect('webapp:update_program', pk=program.pk)


class ExportWord(View):
    def export_program(self, pr):
        document = Document()
        document.add_paragraph('Программа составленна: ' '%s' % str(pr.author))
        document.add_paragraph('Дата скачивания программы: ' "%s" % date.today().strftime('%d. %m. %Y'))
        document.add_paragraph('Ребенок: ' '%s' % str(pr.child))
        document.add_paragraph()

        table = document.add_table(rows=1, cols=4)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Код'
        hdr_cells[1].text = 'Описание навыка'
        hdr_cells[2].text = 'Критерии'
        hdr_cells[3].text = 'Доп.цель'
        document.add_paragraph()
        for i in pr.program_skill.all():
            row_cells = table.add_row()
            row_cells.cells[0].text = str(i.level.skill.category)
            if len(i.goal.all()) > 0:
                for j in i.goal.all():
                    row_cells = table.add_row()
                    row_cells.cells[0].text = str(i.level.skill.code)
                    row_cells.cells[1].text = str(i.level.skill.description)
                    if i.add_creteria:
                        row_cells.cells[2].text = str(i.add_creteria)
                    else:
                        row_cells.cells[2].text = str(i.level.criteria)
                    row_cells.cells[3].text = str(j.goal)
            else:
                row_cells = table.add_row()
                row_cells.cells[0].text = str(i.level.skill.code)
                row_cells.cells[1].text = str(i.level.skill.description)
                row_cells.cells[2].text = str(i.level.criteria)
        return document

    def get(self, request, *args, **kwargs):
        pr = get_object_or_404(Program, pk=self.kwargs.get('pk'))
        document = self.export_program(pr)
        docx_title = pr.name + '   ' + date.today().strftime('%d. %m. %Y') + '.docx'

        f = BytesIO()
        document.save(f)
        length = f.tell()
        f.seek(0)
        response = HttpResponse(
            f.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename=' + docx_title
        response['Content-Length'] = length
        return response


class OpenCloseView(View):
    def get(self, request, **kwargs):
        data = json.loads(request.body)
        goals = get_object_or_404(ProrgamSkillGoal, pk=data['pk'])
        if goals.status == GOAL_STATUS_CLOSED:
            return Http404('Вы пытаетесть  статус закрытого навыка')
        goals.status = data['status']
        goals.save()












