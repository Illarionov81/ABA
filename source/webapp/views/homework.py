import json
from datetime import date
from io import BytesIO

from docx import *
from django.http import JsonResponse, HttpResponse
from django.shortcuts import get_object_or_404, redirect
from django.urls import reverse
from django.views import View
from django.views.generic import ListView, TemplateView, DeleteView
from webapp.models import Program, ProrgamSkillGoal, HomeWork


class HomeworkListView(ListView):
    template_name = 'homework/homeworks_list.html'
    paginate_by = 15
    paginate_orphans = 0
    model = HomeWork

    def get_context_data(self, *, object_list=None, **kwargs):
        context = super().get_context_data(**kwargs)
        program = get_object_or_404(Program, pk=self.kwargs.get('pk'))
        homeworks = HomeWork.objects.filter(program=program)
        context['homeworks'] = homeworks
        context['program'] = program
        return context


class HomeWorkDeleteView(DeleteView):
    template_name = 'homework/homework_delete.html'
    model = HomeWork

    def get_success_url(self):
        return reverse('webapp:homeworks', kwargs={'pk': self.object.program.pk})


class HomeworkCreateView(View):
    def get(self, request, *args, **kwargs):
        program = get_object_or_404(Program, pk=kwargs.get('pk'))
        print(kwargs.get('pk'))
        homework = HomeWork.objects.create(program=program)
        return redirect('webapp:homework_update', pk=homework.pk)


class HomeworkUpdateView(TemplateView):
    template_name = 'homework/homework_create.html'
    model = HomeWork

    def get_context_data(self, *, object_list=None, **kwargs):
        context = super().get_context_data(**kwargs)
        homework = get_object_or_404(HomeWork, pk=self.kwargs.get('pk'))
        goal = ProrgamSkillGoal.objects.filter(skill__program=homework.program.pk)
        goal_in_homework = homework.skill.values_list('pk', flat=True)
        context['goal_in_homework'] = goal_in_homework
        context['goals'] = goal
        context['homework'] = homework
        context['programs'] = homework.program
        return context


class HomeworkAddSkill(View):
    def post(self, request, *args, **kwargs):
        homework = get_object_or_404(HomeWork, pk=kwargs.get('pk'))
        data = json.loads(request.body)
        skill = ProrgamSkillGoal.objects.get(pk=data['id'])
        homework.skill.add(skill)
        return JsonResponse({'add': 'add'})


class HomeworkDeleteSkill(View):
    def delete(self, request, *args, **kwargs):
        homework = get_object_or_404(HomeWork, pk=kwargs.get('pk'))
        data = json.loads(request.body)
        skill = ProrgamSkillGoal.objects.get(pk=data['id'])
        homework.skill.remove(skill)
        return JsonResponse({'remove': 'remove'})


class HomeworkExportWord(View):
    def export_program(self, homework):
        document = Document()
        document.add_paragraph('Домашнее задание: ' '%s' % str(homework.created_date.strftime('%d. %m. %Y')))
        document.add_paragraph('Терапист: ' '%s' % str(homework.program.author))
        document.add_paragraph('Дата скачивания программы: ' "%s" % date.today().strftime('%d. %m. %Y'))
        document.add_paragraph('Ребенок: ' '%s' % str(homework.program.child))
        document.add_paragraph()

        table = document.add_table(rows=1, cols=5)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Код'
        hdr_cells[1].text = 'Описание навыка'
        hdr_cells[2].text = 'Критерии'
        hdr_cells[3].text = 'Доп.цель'
        hdr_cells[4].text = 'Комментарии'
        document.add_paragraph()
        for i in homework.skill.all():
            row_cells = table.add_row()
            row_cells.cells[0].text = str(i.skill.level.skill.category)

            row_cells = table.add_row()
            row_cells.cells[0].text = str(i.skill.level.skill.code)
            row_cells.cells[1].text = str(i.skill.level.skill.description)
            if i.skill.add_creteria:
                row_cells.cells[2].text = str(i.skill.add_creteria)
            else:
                row_cells.cells[2].text = str(i.skill.level.criteria)
            if i.goal:
                row_cells.cells[3].text = str(i.goal)
            else:
                row_cells.cells[3].text = str(i.skill.level.skill.name)
            row_cells.cells[4].text = ''
        return document

    def get(self, request, *args, **kwargs):
        homework = get_object_or_404(HomeWork, pk=self.kwargs.get('pk'))
        document = self.export_program(homework)
        docx_title = homework.program.name + '   ' + date.today().strftime('%d. %m. %Y') + '.docx'

        f = BytesIO()
        document.save(f)
        length = f.tell()
        f.seek(0)
        response = HttpResponse(
            f.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename=' + docx_title
        response['Content-Length'] = length
        return response
